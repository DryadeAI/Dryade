"""Unit tests for document parsers (DOCX, XLSX, HTML).

DOCX and XLSX tests use real fixture files created via python-docx / openpyxl.
HTML tests use plain HTML written to tmp_path files.
"""

from __future__ import annotations

import pytest

# =============================================================================
# DOCX parser tests
# =============================================================================

class TestParseDocx:
    """Tests for parse_docx using real .docx fixture files."""

    @pytest.fixture(autouse=True)
    def _check_docx(self):
        pytest.importorskip("docx")

    def test_parse_docx_paragraphs(self, tmp_path):
        import docx

        from core.knowledge.parsers import parse_docx

        doc = docx.Document()
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("Second paragraph")
        doc.add_paragraph("Third paragraph")
        path = tmp_path / "test.docx"
        doc.save(str(path))

        result = parse_docx(path)
        assert "First paragraph" in result
        assert "Second paragraph" in result
        assert "Third paragraph" in result

    def test_parse_docx_with_table(self, tmp_path):
        import docx

        from core.knowledge.parsers import parse_docx

        doc = docx.Document()
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "A1"
        table.rows[0].cells[1].text = "B1"
        table.rows[1].cells[0].text = "A2"
        table.rows[1].cells[1].text = "B2"
        path = tmp_path / "table.docx"
        doc.save(str(path))

        result = parse_docx(path)
        assert "A1" in result
        assert "B1" in result
        assert "A2" in result
        assert "B2" in result

    def test_parse_docx_empty(self, tmp_path):
        import docx

        from core.knowledge.parsers import parse_docx

        doc = docx.Document()
        path = tmp_path / "empty.docx"
        doc.save(str(path))

        result = parse_docx(path)
        assert result.strip() == ""

# =============================================================================
# XLSX parser tests
# =============================================================================

class TestParseXlsx:
    """Tests for parse_xlsx using real .xlsx fixture files."""

    @pytest.fixture(autouse=True)
    def _check_openpyxl(self):
        pytest.importorskip("openpyxl")

    def test_parse_xlsx_single_sheet(self, tmp_path):
        import openpyxl

        from core.knowledge.parsers import parse_xlsx

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Data"
        ws.append(["Name", "Age"])
        ws.append(["Alice", 30])
        ws.append(["Bob", 25])
        path = tmp_path / "test.xlsx"
        wb.save(str(path))

        result = parse_xlsx(path)
        assert "## Sheet: Data" in result
        assert "Alice" in result
        assert "Bob" in result
        assert "30" in result

    def test_parse_xlsx_multiple_sheets(self, tmp_path):
        import openpyxl

        from core.knowledge.parsers import parse_xlsx

        wb = openpyxl.Workbook()
        ws1 = wb.active
        ws1.title = "Sheet1"
        ws1.append(["Data1"])
        ws2 = wb.create_sheet("Sheet2")
        ws2.append(["Data2"])
        path = tmp_path / "multi.xlsx"
        wb.save(str(path))

        result = parse_xlsx(path)
        assert "## Sheet: Sheet1" in result
        assert "## Sheet: Sheet2" in result
        assert "Data1" in result
        assert "Data2" in result

    def test_parse_xlsx_empty_cells(self, tmp_path):
        import openpyxl

        from core.knowledge.parsers import parse_xlsx

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Sparse"
        ws["A1"] = "Hello"
        # B1 is None
        ws["C1"] = "World"
        path = tmp_path / "sparse.xlsx"
        wb.save(str(path))

        result = parse_xlsx(path)
        assert "None" not in result
        assert "Hello" in result
        assert "World" in result

# =============================================================================
# HTML parser tests
# =============================================================================

class TestParseHtml:
    """Tests for parse_html -- uses stdlib html.parser via BeautifulSoup."""

    @pytest.fixture(autouse=True)
    def _check_bs4(self):
        pytest.importorskip("bs4")

    def test_parse_html_strips_tags(self, tmp_path):
        from core.knowledge.parsers import parse_html

        html = "<html><body><p>Hello</p><div>World</div><span>!</span></body></html>"
        path = tmp_path / "test.html"
        path.write_text(html, encoding="utf-8")

        result = parse_html(path)
        assert "Hello" in result
        assert "World" in result
        assert "<p>" not in result
        assert "<div>" not in result

    def test_parse_html_removes_scripts(self, tmp_path):
        from core.knowledge.parsers import parse_html

        html = """<html><body>
        <script>alert('x')</script>
        <style>.foo { color: red; }</style>
        <p>Content here</p>
        </body></html>"""
        path = tmp_path / "scripts.html"
        path.write_text(html, encoding="utf-8")

        result = parse_html(path)
        assert "alert" not in result
        assert ".foo" not in result
        assert "Content here" in result

    def test_parse_html_removes_nav_footer(self, tmp_path):
        from core.knowledge.parsers import parse_html

        html = """<html><body>
        <header>Site Header</header>
        <nav>Navigation Menu</nav>
        <main><p>Main content</p></main>
        <footer>Copyright 2026</footer>
        </body></html>"""
        path = tmp_path / "nav.html"
        path.write_text(html, encoding="utf-8")

        result = parse_html(path)
        assert "Site Header" not in result
        assert "Navigation Menu" not in result
        assert "Copyright 2026" not in result
        assert "Main content" in result

    def test_parse_html_preserves_meaningful_text(self, tmp_path):
        from core.knowledge.parsers import parse_html

        html = """<html><body>
        <article>
            <h1>Important Title</h1>
            <p>This is the main article content with <strong>bold</strong> text.</p>
        </article>
        </body></html>"""
        path = tmp_path / "article.html"
        path.write_text(html, encoding="utf-8")

        result = parse_html(path)
        assert "Important Title" in result
        assert "main article content" in result
        assert "bold" in result
