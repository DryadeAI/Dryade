"""Document parsers for Knowledge/RAG Pipeline.

Standalone parser functions for DOCX, XLSX, and HTML file formats.
Each returns plain text suitable for chunking and embedding.
"""

from __future__ import annotations

from pathlib import Path


def parse_docx(path: Path) -> str:
    """Parse a DOCX file to plain text.

    Extracts text from paragraphs and tables.

    Args:
        path: Path to the .docx file.

    Returns:
        Extracted plain text content.

    Raises:
        ImportError: If python-docx is not installed.
    """
    try:
        import docx
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX parsing. Install it with: pip install python-docx"
        )

    document = docx.Document(str(path))

    # Extract paragraphs
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

    # Extract tables
    table_lines: list[str] = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_lines.append("\t".join(cells))

    parts = paragraphs
    if table_lines:
        parts = parts + [""] + table_lines

    return "\n".join(parts)

def parse_xlsx(path: Path) -> str:
    """Parse an XLSX file to plain text (sheet-by-sheet, row-by-row).

    Args:
        path: Path to the .xlsx file.

    Returns:
        Extracted plain text content with sheet headers.

    Raises:
        ImportError: If openpyxl is not installed.
    """
    try:
        import openpyxl
    except ImportError:
        raise ImportError(
            "openpyxl is required for XLSX parsing. Install it with: pip install openpyxl"
        )

    wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    parts: list[str] = []

    for sheet in wb.worksheets:
        parts.append(f"## Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(cell) for cell in row if cell is not None]
            if cells:
                parts.append("\t".join(cells))

    wb.close()
    return "\n".join(parts)

def parse_html(path: Path) -> str:
    """Parse an HTML file to plain text.

    Strips scripts, styles, navigation, and other non-content elements.

    Args:
        path: Path to the .html file.

    Returns:
        Extracted meaningful text content.

    Raises:
        ImportError: If beautifulsoup4 is not installed.
    """
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        raise ImportError(
            "beautifulsoup4 is required for HTML parsing. "
            "Install it with: pip install beautifulsoup4"
        )

    html = path.read_text(encoding="utf-8")
    soup = BeautifulSoup(html, "html.parser")

    # Remove non-content tags
    for tag_name in ("script", "style", "nav", "footer", "header"):
        for tag in soup.find_all(tag_name):
            tag.decompose()

    return soup.get_text(separator="\n", strip=True)
